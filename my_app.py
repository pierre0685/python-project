from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

#profile picture
document.add_picture(
    'profile-pic.png', 
    width=Inches(2.0))


#name phone number and email details
name = input("What is your name? ")
speak('Hello ' + name + ' how are you doing today?')

speak('What is your phone number?')
phone_number = input("What is your phone number? ")
email = input("What is your email? ")

document.add_paragraph(
name + ' | ' + phone_number + ' | ' + email)

#about me
document.add_heading('About me')
document.add_paragraph(
    input("Tell me about yourself? "))


#work expereince
document.add_heading("Work Experience")
p = document.add_paragraph()

company = input("Enter company ")
from_date = input("From Date ")
to_date = input("To Date ")

p.add_run(company + ' ').bold = True
p.add_run(from_date + ' -' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experince at  ' + company)
p.add_run(experience_details)

#more experiences
while True:
    has_more_experiences = input(
        'Do you have more experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input("Enter company ")
        from_date = input("From Date ")
        to_date = input("To Date ")

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + ' -' + to_date + '\n').italic = True

        experience_details = input(
        'Describe your experince at ' + company + ' ')
        p.add_run(experience_details)
    else:
        break

#skills
document.add_heading('Skills')
skill = input("Enter Skill ")
p = document.add_paragraph(skill)
p.style = 'List Bullet'

#more skills
while True:
    has_more_skills = input(
        'Do you have more skills? Yes or No ')
    if has_more_skills.lower() =='yes':
        p = document.add_paragraph()

        skill = input("Enter Skill ")
        p = document.add_paragraph(skill)
        p.style = 'List Bullet' 
    else:
        break

#footer
section = document.sections [0]
footer = section.footer
p = footer.paragraphs [0]
p.text = "CV generated using Amigoscode and Intuit Quickbooks course project "
    
document.save('cv.docx')